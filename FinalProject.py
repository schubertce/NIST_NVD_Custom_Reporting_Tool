import docx
import requests
import re
import sys
import pyinputplus as pyip
import smtplib
import ssl
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from bs4 import BeautifulSoup

# Welcome message
print('--------------------------------------------------------------------------')
print('Welcome to the NIST National Vulnerability Database custom reporting tool!')
print('--------------------------------------------------------------------------\n\n')

# Get user input
print('Please select the month and year you wish pull data from for your report.')
print('-------------------------------------------------------------------------')
year = pyip.inputInt('Please input a year between 1996 and 2020: ', min=1996, max=2020)
month = pyip.inputInt('Please input a numerical value for the month from 1 to 12: ', min=1, max=12)
search = pyip.inputYesNo('\nDo you want to search for a specific keyword? Please enter yes or no: ')

# Prompt user to enter a search term if they elected to use the keyword search function
if search == 'yes':
    keywords = input('Keyword: ')
    print('Please wait while your report is being compiled...')
else:
    keywords = ''
    print('Please wait while your report is being compiled...')

# Create report document in Microsoft Word and insert heading
threatReport = docx.Document()
threatReport.add_heading('National Vulnerability Database', 0)
threatReport.add_heading('Custom Report', 0)

# Load main webpage into Python using input from user to specify which month and year
mainURL = 'https://nvd.nist.gov/vuln/full-listing/' + str(year) + '/' + str(month)
mainPage = requests.get(mainURL)

# Parse HTML data
soup = BeautifulSoup(mainPage.content, 'html.parser')

# Assign parsed data to a variable
results = soup.find(id='body-section')

# Search the variable for specific HTML elements that contain URL values and put them in a list.
threats = results.find_all('span', class_='col-md-2')

# Variable to hold the number of keyword hits if the user selects that option.
hits = 0

# Function that searches the description of each scraped vulnerability for
# keyword and then writes the hits to the report or exits program if no hits.
def scraperKeyword(threatDesc, threatName, threatPub, threatSrc, threatURL):
    global hits
    if keywords.lower() in threatDesc.lower():
        threatReport.add_heading(str(threatName), level=2)
        threatReport.add_paragraph(str(threatPub))
        threatReport.add_paragraph(str(threatDesc))
        threatReport.add_paragraph(str(threatStatus))
        threatReport.add_paragraph(str(threatSrc))
        threatReport.add_paragraph(str(threatURL))
        hits += 1
    else:
        pass

# Function that is used if keyword searching is declined and writes all
# vulnerabilities for the month to the report
def scraper(threatDesc, threatName, threatPub, threatSrc, threatURL):
    global hits
    threatReport.add_heading(str(threatName), level=2)
    threatReport.add_paragraph(str(threatPub))
    threatReport.add_paragraph(str(threatDesc))
    threatReport.add_paragraph(str(threatStatus))
    threatReport.add_paragraph(str(threatSrc))
    threatReport.add_paragraph(str(threatURL))
    hits =+ 1

# limits the program to only 5 items in report for the presentation only
index = 0
limit = 25

# Loop to step through list containing URLs
for threat in threats:

    # Use regex to isolate the URL from the HTML code
    threatPage = re.findall('href=[\"\'](.*?)[\"\']', str(threat))

    # Load individual vulnerability webpage into Python
    threatURL = 'https://nvd.nist.gov' + ' '.join(threatPage[:5])
    threatPage = requests.get(threatURL)

    # Parse HTML data
    threatSoup = BeautifulSoup(threatPage.content, 'html.parser')

    # Assign parsed data to variable
    threatResults = threatSoup.find(id='vulnDetailPanel')

    # Search parsed data for specific data points, clean them up, and assign them to string variables
    threatName = re.sub('<span data-testid="page-header-vuln-id">', '', re.sub('</span>', '', str(threatResults.select('span[data-testid="page-header-vuln-id"]')))).strip('[]')
    threatDesc = 'Description: ' + re.sub('u\'', '', re.sub('<p data-testid="vuln-description">', '', re.sub('</p>', '', str(threatResults.select('p[data-testid="vuln-description"]'))))).strip('[]')
    threatPub = 'Published: ' + re.sub('<span data-testid="vuln-published-on">', '', re.sub('</span>', '', str(threatResults.select('span[data-testid="vuln-published-on"]')))).strip('[]')
    threatSrc = 'Source: ' + re.sub('<span data-testid="vuln-current-description-source">', '', re.sub('</span>', '', str(threatResults.select('span[data-testid="vuln-current-description-source"]')))).strip('[]')
    threatStatus = 'Status: ' + re.sub('<span data-testid="vuln-warning-status-name">', '', re.sub('</span>', '', str(threatResults.select('span[data-testid="vuln-warning-status-name"]')))).strip('[]')

    # Once the CVE has been analyzed the status field is removed, so this if statement
    # sets the status as "Analyzed" if there is no status field available.
    if threatStatus == 'Status: ':
        threatStatus = 'Status: Analyzed'
    else:
        pass

    # Calls either of the two functions above depending on if keyword searching is selected
    if search == 'yes':
        scraperKeyword(threatDesc, threatName, threatPub, threatSrc, threatURL)
    else:
        scraper(threatDesc, threatName, threatPub, threatSrc, threatURL)

    # limits the program to only 5 items in report for the presentation only
    index += 1
    if index == limit:
        break

# If no hits were found the program ends, otherwise it prompts user for a path to save report and stores it in variable
if search == 'yes' and hits == 0:
    sys.exit('\nThere were no hits on your keyword and no report was created.')
else:
    if search == 'yes':
        print('There were', str(hits), 'hit(s) on your keyword.')
    else:
        pass
    print('\n')
    path = pyip.inputFilepath(r'Please specify a path and filename to save your report (i.e. C:\FinalProject\Report.docx): ', mustExist=True)

# Saves the report
threatReport.save(path)

# Report complete message
print('\nYour report is complete!\n')

# Asks the user if they would like to email the newly created report.
sendEmail = pyip.inputYesNo('Would you like to email this report? Please enter yes or no: ')

# if statement that either emails the report or ends the program per users choice.
if sendEmail == 'yes':
    email = pyip.inputEmail('Please enter an email address to send this report to: ')

    subject = "NIST NVD Custom Report"
    body = "Attached is a custom report of vulnerabilities from NIST's National Vulnerability Database for " + str(month) + '/' + str(year) + '.'
    bodyKeyword = "Attached is a custom report of vulnerabilities from NIST's National Vulnerability Database for " + str(month) + '/' + str(year) + '. ' + 'This report was culled using the keyword: ' + keywords
    sender_email = 'nvcc.itn270.finalproject@gmail.com'
    receiver_email = email
    password = 'Nor@hjune0511'

    # Create a multipart message and set headers
    message = MIMEMultipart()
    message["From"] = sender_email
    message["To"] = receiver_email
    message["Subject"] = subject
    message["Bcc"] = receiver_email

    # Adds different body sections to the email
    # depending on if keyword searching was selected or not
    if search == 'yes':
        message.attach(MIMEText(bodyKeyword, "plain"))
    else:
        message.attach(MIMEText(body, "plain"))

    filename = path

    # Open PDF file in binary mode
    with open(filename, "rb") as attachment:
        # Add file as application/octet-stream
        # Email client can usually download this automatically as attachment
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())

    # Encode file in ASCII characters to send by email
    encoders.encode_base64(part)

    # Add header as key/value pair to attachment part
    part.add_header("Content-Disposition", f"attachment; filename= {filename}",)

    # Add attachment to message and convert message to string
    message.attach(part)
    text = message.as_string()

    # Log in to server using secure context and send email
    context = ssl.create_default_context()
    with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
        server.login(sender_email, password)
        server.sendmail(sender_email, receiver_email, text)

    print('\nYour report has been sent!\n')
    print('--------------------------------------------------')
    print('Thank you for using the NVD Custom Reporting tool!')
    print('--------------------------------------------------')
else:
    print('\n--------------------------------------------------')
    print('Thank you for using the NVD Custom Reporting tool!')
    print('--------------------------------------------------')
